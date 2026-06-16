"""
cv_extractor.py
----------------
Extracts plain text from PDF, DOCX, and XLSX files LOCALLY (no API call).
This is the key token-saving step: we send only clean text to Claude,
not raw file bytes.
"""

import os
from pathlib import Path


def extract_text_from_pdf(filepath: str) -> str:
    """Extract text from a PDF file."""
    from pypdf import PdfReader
    try:
        reader = PdfReader(filepath)
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
        return "\n".join(text_parts).strip()
    except Exception as e:
        return f"[ERROR reading PDF: {e}]"


def extract_text_from_docx(filepath: str) -> str:
    """Extract text from a Word .docx file."""
    from docx import Document
    try:
        doc = Document(filepath)
        text_parts = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Tables (CVs often have data in tables)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n".join(text_parts).strip()
    except Exception as e:
        return f"[ERROR reading DOCX: {e}]"


def extract_text_from_xlsx(filepath: str) -> str:
    """Extract text from an Excel file."""
    from openpyxl import load_workbook
    try:
        wb = load_workbook(filepath, data_only=True)
        text_parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            text_parts.append(f"--- Sheet: {sheet_name} ---")
            for row in ws.iter_rows(values_only=True):
                row_values = [str(c) for c in row if c is not None and str(c).strip()]
                if row_values:
                    text_parts.append(" | ".join(row_values))
        return "\n".join(text_parts).strip()
    except Exception as e:
        return f"[ERROR reading XLSX: {e}]"


def extract_text(filepath: str) -> str:
    """
    Auto-detect file type and extract text.
    Returns clean text ready to send to Claude.
    """
    ext = Path(filepath).suffix.lower()

    if ext == ".pdf":
        text = extract_text_from_pdf(filepath)
    elif ext in (".docx", ".doc"):
        text = extract_text_from_docx(filepath)
    elif ext in (".xlsx", ".xls"):
        text = extract_text_from_xlsx(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    # Trim excessive whitespace to save tokens
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    cleaned = "\n".join(lines)

    # Safety cap: most CVs fit in ~8000 chars. Hard limit to keep tokens low.
    MAX_CHARS = 15000
    if len(cleaned) > MAX_CHARS:
        cleaned = cleaned[:MAX_CHARS] + "\n[... truncated ...]"

    return cleaned


def get_supported_extensions() -> tuple:
    return (".pdf", ".docx", ".doc", ".xlsx", ".xls")
